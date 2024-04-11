import streamlit as st
from mtranslate import translate
import pandas as pd
import os
import pyautogui
from io import StringIO
import docx

# read language dataset
df = pd.read_excel('language.xlsx')
lang = df['Language'].to_list()
langlist=tuple(lang)
langcode = df['iso'].to_list()

# create dictionary of language and 2 letter langcode
lang_array = {lang[i]: langcode[i] for i in range(len(langcode))}

# layout
st.title("Language Translation app")
st.markdown("Python 🐍 code with Streamlit ! (https://www.streamlit.io/)")
clear = st.button("CLEAR")
inputtext = st.text_area("INPUT",height=200)
choice = st.sidebar.radio('SELECT LANGUAGE',langlist)

#Function for reading a document
def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

uploaded_file = st.file_uploader("Choose a file")
string_data=''
if uploaded_file is not None:
    # To read file as bytes:
    bytes_data = uploaded_file.getvalue()
    #st.write(bytes_data)

    # To convert to a string based IO:
    #stringio = StringIO(uploaded_file.getvalue().decode())
    #st.write(stringio)

    # To read file as string:
    string_data = getText(uploaded_file)
    #st.write(string_data)

# clear function
def Clear():
    pyautogui.press("tab", interval=0.15)
    pyautogui.hotkey("ctrl", "a",'del', interval=0.15)
    pyautogui.press("tab", interval=0.15)

# I/O
if len(string_data) > 0 :
    try:
        output = translate(string_data,lang_array[choice])
        #st.text_area("TRANSLATED TEXT",output,height=200)
        final_doc = docx.Document()
        final_doc.add_paragraph(output)
        final_doc.save('translated_document.docx')
        with open('translated_document.docx', 'rb') as f:
            st.download_button('Download translated document', f, file_name='translated_document.docx')

    except Exception as e:
        st.error(e)

# Clear I/O
if clear:
    Clear()